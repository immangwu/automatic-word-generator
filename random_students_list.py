# import pandas lib as pd
import pandas as pd
import random
import numpy as np
import xlsxwriter
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH 
workbook = xlsxwriter.Workbook('arrays.xlsx')
worksheet = workbook.add_worksheet()
#open the event list and choose the event
data = pd.read_excel('Mech_events_mailmerge.xlsx')
for j in range(len(data.index)):
    E=(data.loc[[j],['Event']]);E=E.values.tolist();E = str(E)[3:-3];
    Da=(data.loc[[j],['Date']]);Da=Da.values.tolist();Da = str(Da)[3:-3];
    NP=(data.loc[[j],['NP']]);NP=NP.values.tolist();NP = str(NP)[2:-2];
    Yr=(data.loc[[j],['year']]);Yr=Yr.values.tolist();Yr = str(Yr)[2:-2];
    text=str(Yr)+'.xlsx'
#based on the year create student database  
    # read by default 1st sheet of an excel file
    dataframe1 = pd.read_excel(text)
    df=dataframe1
    arr = df.to_numpy()
    number_of_rows = arr.shape[0]
    #input_a = int(input("enter the participants"))
    random_indices = np.random.choice(number_of_rows,size=int(NP),replace=False)
    print(random_indices)
    print(random_indices[1])
    #create excel sheet
    a = np.zeros((len(random_indices),3),dtype=object)
    for i in range(0,len(random_indices)):
        a[i,0]=arr[random_indices[i],0];
        a[i,1]=arr[random_indices[i],1];
        #str(a[i,2]);
        a[i,2]=(arr[random_indices[i],2]);
    
    #a[1,i]=arr[random_indices[i],1];
    #a[2,i]=arr[random_indices[i],2];
    #a=np.transpose(a)
    ##col=3
    ##for row, data in enumerate(a):
    ##    worksheet.write_column(row, col, data)
    ##
    ##print(a)
    ##    
    ##workbook.close()
    print(a)
    df = pd.DataFrame(a)
    df.to_excel(excel_writer = "test.xlsx")
#Content for word file Event details

    da=pd.read_excel('test.xlsx')
#create word file



    document = Document()
    document.add_picture('srit.png', width=Inches(1.25))
    document.add_heading(str(E), 0)
    sen='Attendence sheet for the event '+str(E)
    p = document.add_paragraph(sen)
    p.alignment = 3 
    p.add_run(' conducted on ')
    p.add_run(str(Da)).bold = True
    p.add_run(' Organized by ')
    p.add_run(' Department of Mechanical Engineering').italic = True

    document.add_heading('Participants Details', level=1).underline = True
    #document.add_paragraph('', style='Intense Quote')

    #document.add_paragraph(
    #    'first item in unordered list', style='List Bullet'
    #)
    #document.add_paragraph(
    #    'first item in ordered list', style='List Number'
    #)

    #document.add_picture('srit.png', width=Inches(1.25))

    #records = (
    #    (3, '101', 'Spam'),
    #    (7, '422', 'Eggs'),
    #    (4, '631', 'Spam, spam, eggs, and spam')
    #)
    table = document.add_table(rows=1, cols=3)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'S.No'
    hdr_cells[1].text = 'Register Number'
    hdr_cells[2].text = 'Name'
    for i in range(len(da.index)):
        RN=(da.loc[[i],[1]]);RN=RN.values.tolist();RN = str(RN)[2:-2];#Register number
        Na=(da.loc[[i],[2]]);Na=Na.values.tolist();Na = str(Na)[3:-3];#Name
        
        #for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(i+1)
        row_cells[1].text = RN
        row_cells[2].text = str(Na)
    
    document.add_paragraph('', style='Intense Quote')
    sen='This is to certify that the above students were participated in the event titled '+str(E)
    p = document.add_paragraph(sen)
    p.alignment = 3 
    p.add_run(' conducted on ')
    p.add_run(str(Da)).bold = True
    p.add_run(' Organized by ')
    p.add_run(' Department of Mechanical Engineering').italic = True
    
    paragraph = document.add_paragraph("                                                          ")
    paragraph = document.add_paragraph("Faculty Coordinator")
    paragraph.alignment = 2 # for left, 1 for center, 2 right, 3 justify ....
    document.add_paragraph('', style='Intense Quote')
    document.add_page_break()
    tex=str(E)+'.docx'
    document.save(tex)
